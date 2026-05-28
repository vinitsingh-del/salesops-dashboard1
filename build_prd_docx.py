from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "SalesOps_Dashboard_PRD.docx"


COLORS = {
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "ink": RGBColor(34, 34, 34),
    "muted": RGBColor(90, 90, 90),
    "light_fill": "F2F4F7",
    "blue_fill": "E8EEF5",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_run(paragraph, text, bold=False, italic=False, size=None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.style = "Body Text"
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text)
    p.style = "List Bullet"
    return p


def add_number(doc, text):
    p = doc.add_paragraph(text)
    p.style = "List Number"
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, COLORS["light_fill"])
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(header)
        run.bold = True
        run.font.color.rgb = COLORS["dark_blue"]
        run.font.size = Pt(9.5)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            paragraph = row.cells[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value))
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLORS["ink"]
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["blue_fill"])
    set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
    p = cell.paragraphs[0]
    add_run(p, title + ": ", bold=True, color=COLORS["dark_blue"])
    add_run(p, body, color=COLORS["ink"])
    set_table_width(table, [6.25])
    doc.add_paragraph()


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = COLORS["ink"]
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    body = doc.styles["Body Text"]
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.font.color.rgb = COLORS["ink"]
    body.paragraph_format.space_after = Pt(6)
    body.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 16, 8),
        ("Heading 2", 13, COLORS["blue"], 12, 6),
        ("Heading 3", 12, COLORS["dark_blue"], 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.10


def add_cover(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(6)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_run(title, "SalesOps Dashboard", bold=True, size=28, color=COLORS["dark_blue"])

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    add_run(subtitle, "Product Requirements Document", size=16, color=COLORS["muted"])

    add_callout(
        doc,
        "Purpose",
        "Define the operating logic, user workflows, backend requirements, and failure-audit rules for the SalesOps brand pipeline dashboard.",
    )

    metadata = [
        ("Product", "SalesOps Brand Pipeline Dashboard"),
        ("Primary Users", "Sales, CAM, Finance, Superadmin / Leadership"),
        ("Backend", "Supabase deals table and deal-documents storage bucket"),
        ("Hosting", "GitHub Pages"),
        ("Version", "1.0"),
    ]
    add_table(doc, ["Field", "Detail"], metadata, [1.7, 4.6])
    doc.add_page_break()


def build_doc():
    doc = Document()
    configure_styles(doc)
    add_cover(doc)

    add_heading(doc, "1. Product Summary", 1)
    add_body(
        doc,
        "SalesOps Dashboard is a lightweight web dashboard for managing brand deals from lead creation through agreement, invoicing, onboarding, active delivery, renewal, churn risk, and churn.",
    )
    add_body(
        doc,
        "The product gives sales, CAM, finance, and leadership teams a shared source of truth for deal status, document completeness, invoice readiness, payment follow-up, renewal tracking, and operational failure alerts.",
    )

    add_heading(doc, "2. Goals", 1)
    for item in [
        "Create one place to create, update, review, and delete deals.",
        "Make commercial status clear across approval, agreement, invoice raised, invoice paid, onboarding, active, renewal due, and churned.",
        "Separate first invoice logic from recurring MRR invoice logic.",
        "Track agreement, NDA, and Refrens invoice documents with upload and download options.",
        "Surface churn possibility as Low, Medium, or High before a brand actually churns.",
        "Automatically flag system failure cases that would break sales, finance, or onboarding workflows.",
        "Keep the interface usable on desktop and mobile.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3. Primary Users", 1)
    user_rows = [
        ("Sales / BDM", "Creates new deals, tracks approvals, follows up on agreement and invoice milestones."),
        ("CAM", "Tracks onboarding, active delivery, next action, churn possibility, and renewal readiness."),
        ("Finance", "Checks invoice raised, invoice paid, due date, clearance date, and Refrens invoice copies."),
        ("Superadmin / Leadership", "Reviews all deals, document completeness, churn exposure, failure audit, and pipeline health."),
    ]
    add_table(doc, ["User", "Needs"], user_rows, [1.8, 4.5])

    add_heading(doc, "4. Core Workflows", 1)
    add_heading(doc, "New Lead to Active Deal", 2)
    for step in [
        "Sales creates a new deal from + Deal.",
        "Sales enters brand name, value, deal type, duration, month, owners, and next action date.",
        "Sales marks agreement status and uploads or pastes the agreement copy if available.",
        "Sales raises first invoice and adds Refrens invoice number or uploads the invoice copy.",
        "Finance marks invoice as paid only after payment is received.",
        "Finance adds payment clearance date.",
        "CAM completes onboarding.",
        "MRR deal becomes Active only when agreement is complete and first invoice is paid.",
    ]:
        add_number(doc, step)

    add_heading(doc, "Existing Lead / Existing Deal", 2)
    for step in [
        "Team searches brand in All Deals.",
        "Team updates stage, invoice status, next action, churn possibility, or documents.",
        "If documents are missing, Document Vault remains visible until required files are uploaded.",
        "If a deal has broken logic, Rules > System Failure Audit shows the failure and correction step.",
    ]:
        add_number(doc, step)

    add_heading(doc, "Churn Risk Deal", 2)
    for step in [
        "Team marks churn possibility as Low, Medium, or High.",
        "Medium or High churn requires a churn reason or note.",
        "High churn is shown as operational risk.",
        "Team records next action and owner follow-up.",
        "If the brand exits, it is moved to churn tracking.",
    ]:
        add_number(doc, step)

    add_heading(doc, "5. Functional Requirements", 1)
    add_heading(doc, "Deal Management", 2)
    for item in [
        "Create, update, and delete deals.",
        "Edit deal stage, invoice status, churn possibility, next action, and action date.",
        "Search deals by brand, owner, CAM, or Refrens ID.",
        "Filter by month, quarter, stage, MRR, one-time, and churn risk.",
    ]:
        add_bullet(doc, item)

    field_rows = [
        ("Commercial", "Brand name, deal value, deal type, deal duration, month, stage, Refrens ID."),
        ("Ownership", "Sales owner, CAM owner, business owner / client POC."),
        ("Agreement", "Agreement signed, agreement sent date, signed date, agreement link/copy."),
        ("Invoice", "First invoice raised, invoice number, invoice date, payment terms, due date, clearance date."),
        ("Renewal", "Renewal date, next action, next action date."),
        ("Churn", "Churn possibility and churn notes."),
    ]
    add_table(doc, ["Field Group", "Fields"], field_rows, [1.7, 4.6])

    add_heading(doc, "Document Management", 2)
    for item in [
        "Upload NDA, agreement copy, and Refrens invoice copy.",
        "Store uploaded files in Supabase Storage.",
        "Show uploaded document state in Document Vault.",
        "Allow Superadmin to open or download uploaded documents.",
        "Allow replacement upload if a document is missing or incorrect.",
        "Flag Supabase upload failures.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Invoice Logic", 2)
    invoice_rows = [
        ("One-Time", "Creates one invoice flow only."),
        ("MRR", "Creates recurring invoice schedule based on duration and invoice day."),
        ("Partial Invoice", "First invoice amount may differ from recurring invoice amount."),
        ("Paid Gate", "Paid status requires invoice raised and payment clearance date."),
        ("Overdue Gate", "Pending invoices with past due dates are flagged as overdue."),
    ]
    add_table(doc, ["Logic Area", "Requirement"], invoice_rows, [1.6, 4.7])

    add_heading(doc, "6. Dashboard Views", 1)
    view_rows = [
        ("Dashboard", "Pipeline value, invoices raised, collected value, renewals due, churned, high churn risk, agreement pending, no next action date, funnel, active deals, renewals, churns."),
        ("All Deals", "Searchable deal table with stage, invoice, churn marker, delete action, and journey modal."),
        ("Invoices", "Invoice status by deal, first invoice, recurring invoices, due dates, and payment status."),
        ("Renewals", "Upcoming renewals, urgency by date, owner actions."),
        ("Churned", "Churned brands, churn reason, last invoice status, recovery possibility."),
        ("Documents", "NDA, agreement, Refrens invoice state, upload option, and download option."),
        ("Owners", "Sales owner and CAM visibility."),
        ("Rules", "Operating rules, System Failure Audit, audit coverage checklist, broken flow items."),
    ]
    add_table(doc, ["View", "Purpose"], view_rows, [1.35, 4.95])

    add_heading(doc, "7. System Failure Audit", 1)
    add_body(doc, "The dashboard must automatically flag the following failure cases and show brand, severity, correction step, owner, next action, and action date.")
    failure_rows = [
        ("No next action date", "High", "Add next action date and owner follow-up."),
        ("Agreement missing", "High", "Capture agreement status before invoice or active stages."),
        ("Agreement copy missing", "High", "Upload agreement copy or paste agreement link."),
        ("Invoice paid but invoice not raised", "High", "Add Refrens invoice number/copy before paid status."),
        ("Invoice overdue", "High", "Escalate collections and update next action."),
        ("Payment marked paid but clearance date missing", "High", "Add payment clearance date."),
        ("First invoice raised before agreement sent", "High", "Fix sequence or add agreement sent date."),
        ("MRR duration missing", "Medium", "Set duration so renewal and invoice logic can work."),
        ("MRR invoice schedule missing", "Medium", "Regenerate recurring invoice schedule."),
        ("Churn Medium/High without reason", "Medium", "Add churn reason or notes."),
        ("Supabase document upload failed", "High", "Check storage bucket/policy and re-upload."),
    ]
    add_table(doc, ["Audit Check", "Severity", "Correction"], failure_rows, [2.55, 0.75, 3.0])

    add_heading(doc, "8. Backend Requirements", 1)
    backend_rows = [
        ("Deals table", "public.deals"),
        ("Primary key", "deal_id text"),
        ("Payload", "data jsonb"),
        ("Timestamps", "created_at, updated_at"),
        ("Storage bucket", "deal-documents"),
        ("Stored document types", "nda, agreementCopy, refrensInvoice"),
    ]
    add_table(doc, ["Component", "Requirement"], backend_rows, [2.0, 4.3])
    add_callout(
        doc,
        "Security note",
        "The current lightweight deployment uses public anon access policies for demo/team use. Production should add Supabase Auth and role-based access before sensitive client documents are stored.",
    )

    add_heading(doc, "9. Non-Functional Requirements", 1)
    for item in [
        "Mobile responsive layout.",
        "Fast loading on GitHub Pages.",
        "No build step required for GitHub Pages deployment.",
        "Clear visual status badges.",
        "Low-friction document upload.",
        "Clear error messages when Supabase table or storage setup is missing.",
        "Avoid hidden failure states.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "10. Success Metrics", 1)
    metric_rows = [
        ("Next actions", "100% of active deals have a next action date."),
        ("Paid invoices", "100% of paid deals have invoice raised and clearance date."),
        ("Agreements", "100% of signed agreements have linked/uploaded agreement copy."),
        ("MRR readiness", "0 MRR deals without duration or invoice schedule."),
        ("Churn risk", "0 Medium/High churn risk deals without reason."),
        ("Team efficiency", "Reduced manual follow-up in finance and CAM reviews."),
    ]
    add_table(doc, ["Metric", "Target"], metric_rows, [1.7, 4.6])

    add_heading(doc, "11. Known Limitations", 1)
    for item in [
        "Refrens invoice number is inferred from uploaded file names; true Refrens API sync is not implemented.",
        "Public anon Supabase access is acceptable only for controlled demo/internal use.",
        "GitHub Pages may take a short delay to reflect pushed changes.",
        "Browser file preview depends on Supabase signed URL generation.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "12. Future Enhancements", 1)
    for item in [
        "Supabase Auth with Superadmin, Sales, CAM, and Finance roles.",
        "Refrens API integration.",
        "Automated invoice reminders.",
        "Email or WhatsApp reminders for next action dates.",
        "CSV export.",
        "Audit history per deal.",
        "Dedicated churn recovery pipeline.",
        "SLA dashboards for agreement, invoice, payment, onboarding, and renewal.",
    ]:
        add_bullet(doc, item)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer, "SalesOps Dashboard PRD", size=9, color=COLORS["muted"])

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
